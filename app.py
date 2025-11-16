from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime
import anthropic
from matieres_data_complete import MATIERES_DATA_COMPLETE, CONTEXTES_MONDIAUX
from database import init_db, save_session, save_unit, get_units_by_teacher

app = Flask(__name__)
app.config['GENERATED_UNITS_FOLDER'] = 'generated_units'

# Créer le dossier generated_units s'il n'existe pas
os.makedirs(app.config['GENERATED_UNITS_FOLDER'], exist_ok=True)

# Initialiser la base de données
init_db()

# Fonction pour expandre les années PEI
def expand_pei_years(data):
    """Expands pei1-2 to pei1 and pei2, pei3-4 to pei3 and pei4"""
    expanded = {}
    for key, value in data.items():
        if key == "pei1-2":
            expanded["pei1"] = value
            expanded["pei2"] = value
        elif key == "pei3-4":
            expanded["pei3"] = value
            expanded["pei4"] = value
        else:
            expanded[key] = value
    return expanded

# Charger et expandre les données des matières
MATIERES_DATA = {}
for matiere_key, matiere_data in MATIERES_DATA_COMPLETE.items():
    MATIERES_DATA[matiere_key] = {
        **matiere_data,
        "objectifs": expand_pei_years(matiere_data.get("objectifs", {}))
    }

print(f"[DEBUG] Années PEI disponibles: {list(MATIERES_DATA['design']['objectifs'].keys())}")

# Ancienne définition (supprimée car remplacée par l'import)
OLD_MATIERES_DATA = {
    "design": {
        "nom": "Design",
        "concepts_cles": ["Communautés", "Communication", "Développement", "Systèmes"],
        "concepts_connexes": ["Adaptation", "Collaboration", "Durabilité", "Ergonomie", "Évaluation", 
                              "Fonction", "Forme", "Innovation", "Invention", "Marchés et tendances", 
                              "Perspective", "Ressources"],
        "objectifs": {
            "pei1-2": {
                "A": ["i. Expliquer et justifier le besoin d'apporter une solution à un problème.",
                      "ii. Indiquer et hiérarchiser les grandes étapes des recherches.",
                      "iii. Décrire les caractéristiques principales d'un produit existant.",
                      "iv. Présenter les principales conclusions des recherches."],
                "B": ["i. Développer une liste des critères de réussite.",
                      "ii. Présenter des idées de conception réalisables.",
                      "iii. Présenter la conception retenue.",
                      "iv. Créer un dessin ou un schéma de planification."],
                "C": ["i. Résumer un plan de travail.",
                      "ii. Démontrer des compétences techniques.",
                      "iii. Suivre le plan pour créer la solution.",
                      "iv. Énumérer les changements apportés."],
                "D": ["i. Résumer des méthodes d'essai.",
                      "ii. Résumer la réussite de la solution.",
                      "iii. Résumer en quoi la solution pourrait être améliorée.",
                      "iv. Résumer les effets de la solution."]
            },
            "pei3-4": {
                "A": ["i. Expliquer et justifier le besoin d'apporter une solution à un problème.",
                      "ii. Construire un plan de recherche qui indique et hiérarchise les recherches.",
                      "iii. Analyser un groupe de produits similaires.",
                      "iv. Développer un énoncé de projet qui présente l'analyse des recherches."],
                "B": ["i. Développer un cahier des charges résumant les critères de réussite.",
                      "ii. Présenter un éventail d'idées de conception réalisables.",
                      "iii. Présenter la conception retenue et résumer les raisons du choix.",
                      "iv. Développer des dessins ou des schémas de planification précis."],
                "C": ["i. Construire un plan logique résumant l'utilisation du temps et des ressources.",
                      "ii. Démontrer des compétences techniques excellentes.",
                      "iii. Suivre le plan afin de créer la solution.",
                      "iv. Expliquer les changements apportés."],
                "D": ["i. Décrire des méthodes d'essai détaillées et pertinentes.",
                      "ii. Expliquer dans quelle mesure la solution est une réussite.",
                      "iii. Décrire en quoi la solution pourrait être améliorée.",
                      "iv. Décrire les effets de la solution."]
            },
            "pei5": {
                "A": ["i. Expliquer et justifier le besoin d'apporter une solution pour un client/public spécifique.",
                      "ii. Identifier et hiérarchiser les recherches primaires et secondaires.",
                      "iii. Analyser une gamme de produits existants.",
                      "iv. Développer un énoncé de projet détaillé."],
                "B": ["i. Développer un cahier des charges énonçant clairement les critères de réussite.",
                      "ii. Développer un éventail d'idées de conception réalisables.",
                      "iii. Présenter la conception retenue et justifier son choix.",
                      "iv. Développer des dessins ou des schémas de planification précis et détaillés."],
                "C": ["i. Construire un plan logique décrivant une utilisation efficace du temps et des ressources.",
                      "ii. Démontrer des compétences techniques excellentes.",
                      "iii. Suivre le plan afin de créer la solution.",
                      "iv. Justifier pleinement les changements apportés."],
                "D": ["i. Élaborer des méthodes d'essai détaillées et pertinentes.",
                      "ii. Évaluer de manière critique la réussite de la solution.",
                      "iii. Expliquer en quoi la solution pourrait être améliorée.",
                      "iv. Expliquer les effets de la solution sur le client ou le public."]
            }
        }
    },
    "langue_litterature": {
        "nom": "Langue et littérature",
        "concepts_cles": ["Communication", "Créativité", "Liens", "Perspective"],
        "concepts_connexes": ["But", "Cadre", "Contexte", "Expression personnelle", "Genre", 
                              "Interpellation du destinataire", "Intertextualité", "Personnage", 
                              "Point de vue", "Structure", "Style", "Thème"],
        "objectifs": {
            "pei1-2": {
                "A": ["i. Identifier et commenter les aspects significatifs des textes.",
                      "ii. Identifier et commenter les choix de l'auteur.",
                      "iii. Justifier ses opinions et idées avec des exemples et explications.",
                      "iv. Identifier les similarités et différences dans et entre les textes."],
                "B": ["i. Employer des structures pour organiser le contenu.",
                      "ii. Organiser ses opinions et idées avec logique.",
                      "iii. Utiliser les outils de présentation et de références."],
                "C": ["i. Produire des textes qui démontrent réflexion et imagination.",
                      "ii. Faire des choix stylistiques (procédés linguistiques, littéraires et visuels).",
                      "iii. Choisir des détails et exemples pertinents."],
                "D": ["i. Utiliser un vocabulaire, des tournures de phrases et des formes d'expression appropriés.",
                      "ii. S'exprimer à l'oral et à l'écrit avec un registre et un style appropriés.",
                      "iii. Utiliser une grammaire, syntaxe et ponctuation correctes.",
                      "iv. Utiliser une orthographe et une prononciation correctes.",
                      "v. Utiliser des techniques de communication non verbale."]
            },
            "pei3-4": {
                "A": ["i. Identifier et expliquer le contenu, le contexte, la langue, la structure, la technique, le style et la relation entre les textes.",
                      "ii. Identifier et expliquer les effets des choix de l'auteur.",
                      "iii. Justifier ses opinions et idées avec des exemples, explications et la terminologie appropriée.",
                      "iv. Interpréter les similarités et différences dans et entre les genres et les textes."],
                "B": ["i. Employer des structures adaptées pour organiser le contenu.",
                      "ii. Organiser ses opinions et idées avec cohérence et logique.",
                      "iii. Utiliser les outils de présentation et de références de manière adaptée."],
                "C": ["i. Produire des textes qui démontrent réflexion, imagination et sensibilité.",
                      "ii. Faire des choix stylistiques en démontrant une conscience des effets sur le public.",
                      "iii. Choisir des détails et exemples pertinents pour développer ses idées."],
                "D": ["i. Utiliser un vocabulaire, tournures et formes d'expression appropriés et variés.",
                      "ii. S'exprimer à l'oral et à l'écrit avec un registre et style appropriés.",
                      "iii. Utiliser une grammaire, syntaxe et ponctuation correctes.",
                      "iv. Utiliser une orthographe et prononciation correctes.",
                      "v. Utiliser des techniques de communication non verbale appropriées."]
            },
            "pei5": {
                "A": ["i. Analyser le contenu, le contexte, la langue, la structure, la technique, le style de chaque texte, et la relation entre les textes.",
                      "ii. Analyser les effets des choix de l'auteur sur son public.",
                      "iii. Justifier ses opinions et idées avec des exemples, explications et terminologie appropriée.",
                      "iv. Évaluer les similarités et différences en associant des caractéristiques dans et entre les genres et les textes."],
                "B": ["i. Employer des structures permettant d'organiser le contenu et convenant au contexte et à l'intention.",
                      "ii. Organiser ses opinions et idées avec constance, cohérence et logique.",
                      "iii. Utiliser les outils de présentation et de références pour parvenir à un style adapté."],
                "C": ["i. Produire des textes qui démontrent perspicacité, imagination et sensibilité, et une réflexion critique.",
                      "ii. Faire des choix stylistiques en démontrant une conscience des effets sur le public.",
                      "iii. Choisir des détails et exemples pertinents pour développer ses idées."],
                "D": ["i. Utiliser un vocabulaire, tournures et formes d'expression appropriés et variés.",
                      "ii. S'exprimer à l'oral et à l'écrit avec un registre et style convenant au contexte et à l'intention.",
                      "iii. Utiliser une grammaire, syntaxe et ponctuation correctes.",
                      "iv. Utiliser une orthographe et prononciation correctes.",
                      "v. Utiliser des techniques de communication non verbale appropriées."]
            }
        }
    }
}

# CONTEXTES_MONDIAUX déjà importé de matieres_data_complete

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/matieres')
def get_matieres():
    """Retourne la liste des matières disponibles"""
    matieres = {key: {"nom": value["nom"]} for key, value in MATIERES_DATA.items()}
    return jsonify(matieres)

@app.route('/api/matiere/<matiere_id>')
def get_matiere_details(matiere_id):
    """Retourne les détails d'une matière"""
    if matiere_id in MATIERES_DATA:
        return jsonify(MATIERES_DATA[matiere_id])
    return jsonify({"error": "Matière non trouvée"}), 404

@app.route('/api/sessions/<enseignant>')
def get_teacher_sessions(enseignant):
    """Récupère les sessions d'un enseignant"""
    try:
        from database import get_recent_sessions
        sessions = get_recent_sessions(limit=20)
        # Filtrer par enseignant
        teacher_sessions = [s for s in sessions if s.get('enseignant') == enseignant]
        return jsonify({"sessions": teacher_sessions})
    except Exception as e:
        print(f"[ERROR] Failed to get sessions: {e}")
        return jsonify({"error": str(e), "sessions": []}), 500

@app.route('/api/session/<session_id>')
def get_session_details(session_id):
    """Récupère les détails d'une session"""
    try:
        from database import get_session_by_id
        session = get_session_by_id(session_id)
        if session:
            return jsonify(session)
        return jsonify({"error": "Session non trouvée"}), 404
    except Exception as e:
        print(f"[ERROR] Failed to get session: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-units', methods=['POST'])
def generate_units():
    """Génère les unités PEI basées sur les chapitres fournis"""
    try:
        # Try to get JSON data, handle if it's None
        data = request.get_json(force=True, silent=True)
        if data is None:
            data = request.json
        
        if data is None:
            return jsonify({"error": "No JSON data received"}), 400
            
        print(f"[DEBUG] Received data: {data}")
        
        matiere_id = data.get('matiere')
        annee_pei = data.get('annee_pei')
        enseignant = data.get('enseignant')
        chapitres = data.get('chapitres', [])
        
        print(f"[DEBUG] matiere_id={matiere_id}, annee_pei={annee_pei}, enseignant={enseignant}")
        print(f"[DEBUG] chapitres count: {len(chapitres)}")
        
        if not matiere_id or not annee_pei or not chapitres:
            return jsonify({"error": "Données manquantes"}), 400
        
        matiere_data = MATIERES_DATA.get(matiere_id)
        if not matiere_data:
            return jsonify({"error": f"Matière non trouvée: {matiere_id}"}), 404
        
        print(f"[DEBUG] Matiere data loaded: {matiere_data.get('nom')}")
        
        # Déterminer le nombre d'unités à générer
        nb_unites = 6 if matiere_id == "langue_litterature" else 4
        print(f"[DEBUG] Generating {nb_unites} units")
        
        # Utiliser l'IA pour regrouper les chapitres et générer les unités
        units = generate_units_with_ai(chapitres, matiere_data, annee_pei, nb_unites, enseignant)
        print(f"[DEBUG] Generated {len(units)} units successfully")
        
        # Sauvegarder la session dans MongoDB
        try:
            session_id = save_session(enseignant, matiere_id, annee_pei, chapitres, units)
            if session_id:
                print(f"[DEBUG] Session saved to MongoDB: {session_id}")
        except Exception as e:
            print(f"[WARNING] Failed to save session to MongoDB: {e}")
        
        return jsonify({"units": units})
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[ERROR] Exception in generate_units: {str(e)}")
        print(f"[ERROR] Traceback: {error_details}")
        return jsonify({"error": str(e), "details": error_details}), 500

def generate_units_with_ai(chapitres, matiere_data, annee_pei, nb_unites, enseignant):
    """Utilise l'IA pour regrouper les chapitres et générer les unités"""
    
    # Obtenir la clé API Claude (si disponible)
    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    
    if api_key:
        try:
            client = anthropic.Anthropic(api_key=api_key)
            
            prompt = f"""Tu es un expert en pédagogie du Programme d'Éducation Intermédiaire (PEI) de l'IB.

Voici les chapitres d'un programme pour la matière "{matiere_data['nom']}" en année {annee_pei}:

{json.dumps(chapitres, indent=2, ensure_ascii=False)}

Voici les concepts clés disponibles: {', '.join(matiere_data['concepts_cles'])}
Voici les concepts connexes disponibles: {', '.join(matiere_data['concepts_connexes'])}
Voici les contextes mondiaux disponibles: {', '.join(CONTEXTES_MONDIAUX)}

Tu dois générer EXACTEMENT {nb_unites} unités pédagogiques en regroupant les chapitres par thèmes cohérents.

Pour chaque unité, fournis:
1. titre_unite: Un titre engageant qui regroupe les chapitres
2. chapitres_inclus: Les numéros/identifiants des chapitres regroupés
3. duree: Durée totale en heures (somme des durées des chapitres)
4. concept_cle: Un concept clé parmi ceux disponibles
5. concepts_connexes: 2-3 concepts connexes pertinents
6. contexte_mondial: Un contexte mondial pertinent
7. enonce_recherche: Un énoncé de recherche stimulant
8. questions_factuelles: 2-3 questions factuelles
9. questions_conceptuelles: 2-3 questions conceptuelles
10. questions_debat: 2-3 questions invitant au débat
11. objectifs_specifiques: Les objectifs spécifiques pertinents (format: "A.i, A.ii, B.i, C.iii, D.ii")

Réponds UNIQUEMENT en JSON valide avec ce format:
{{
  "unites": [
    {{
      "titre_unite": "...",
      "chapitres_inclus": [...],
      "duree": 0,
      "concept_cle": "...",
      "concepts_connexes": ["...", "..."],
      "contexte_mondial": "...",
      "enonce_recherche": "...",
      "questions_factuelles": ["...", "..."],
      "questions_conceptuelles": ["...", "..."],
      "questions_debat": ["...", "..."],
      "objectifs_specifiques": ["A.i", "A.ii", "B.i", ...]
    }}
  ]
}}"""

            message = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}]
            )
            
            response_text = message.content[0].text
            # Nettoyer la réponse si elle contient des balises markdown
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            result = json.loads(response_text)
            return result.get('unites', [])
        
        except Exception as e:
            print(f"Erreur lors de l'utilisation de l'IA: {e}")
            # Fallback vers la génération basique
            pass
    
    # Génération basique si pas d'API ou erreur
    return generate_units_basic(chapitres, matiere_data, annee_pei, nb_unites)

def generate_units_basic(chapitres, matiere_data, annee_pei, nb_unites):
    """Génération basique des unités sans IA"""
    print(f"[DEBUG] generate_units_basic called with {len(chapitres)} chapitres, {nb_unites} units to generate")
    units = []
    chapitres_per_unit = max(1, len(chapitres) // nb_unites)
    
    for i in range(nb_unites):
        start_idx = i * chapitres_per_unit
        end_idx = start_idx + chapitres_per_unit if i < nb_unites - 1 else len(chapitres)
        unit_chapitres = chapitres[start_idx:end_idx]
        
        print(f"[DEBUG] Unit {i+1}: chapitres from {start_idx} to {end_idx}")
        
        duree_totale = sum(ch.get('duree', 0) for ch in unit_chapitres)
        titres = [ch.get('titre', '') for ch in unit_chapitres]
        
        unit = {
            "titre_unite": f"Unité {i+1}: {' et '.join(titres[:2])}",
            "chapitres_inclus": [ch.get('id', start_idx + j) for j, ch in enumerate(unit_chapitres)],
            "duree": duree_totale,
            "concept_cle": matiere_data['concepts_cles'][i % len(matiere_data['concepts_cles'])],
            "concepts_connexes": matiere_data['concepts_connexes'][:3],
            "contexte_mondial": CONTEXTES_MONDIAUX[i % len(CONTEXTES_MONDIAUX)],
            "enonce_recherche": f"Exploration de {titres[0]}",
            "questions_factuelles": [
                f"Quels sont les éléments clés de {titres[0]}?",
                f"Comment définir {titres[0]}?"
            ],
            "questions_conceptuelles": [
                f"Pourquoi {titres[0]} est-il important?",
                f"Comment {titres[0]} influence-t-il notre compréhension?"
            ],
            "questions_debat": [
                f"Dans quelle mesure {titres[0]} affecte-t-il notre société?",
                f"Quel est l'impact de {titres[0]} sur notre futur?"
            ],
            "objectifs_specifiques": ["A.i", "A.ii", "B.i", "C.i", "D.i"]
        }
        units.append(unit)
    
    return units

@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    """Génère le document Word pour une unité"""
    try:
        # Try to get JSON data
        data = request.get_json(force=True, silent=True)
        if data is None:
            data = request.json
        
        if data is None:
            return jsonify({"error": "No JSON data received"}), 400
        
        print(f"[DEBUG] generate_document - Received data keys: {data.keys()}")
        
        unite = data.get('unite')
        matiere_id = data.get('matiere')
        annee_pei = data.get('annee_pei')
        enseignant = data.get('enseignant', '')
        
        print(f"[DEBUG] matiere_id={matiere_id}, annee_pei={annee_pei}, enseignant={enseignant}")
        print(f"[DEBUG] unite titre: {unite.get('titre_unite') if unite else 'None'}")
        
        if not unite or not matiere_id or not annee_pei:
            return jsonify({"error": "Données manquantes"}), 400
        
        matiere_data = MATIERES_DATA.get(matiere_id)
        if not matiere_data:
            return jsonify({"error": f"Matière non trouvée: {matiere_id}"}), 404
        
        print(f"[DEBUG] Starting document generation for: {unite.get('titre_unite')}")
        
        # Générer le document Word
        filename = create_word_document(unite, matiere_data, annee_pei, enseignant)
        
        print(f"[DEBUG] Document generated successfully: {filename}")
        
        return jsonify({"filename": filename, "download_url": f"/download/{filename}"})
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[ERROR] Exception in generate_document: {str(e)}")
        print(f"[ERROR] Traceback: {error_details}")
        return jsonify({"error": str(e), "details": error_details}), 500

def create_word_document(unite, matiere_data, annee_pei, enseignant):
    """Crée le document Word basé sur le template"""
    
    print(f"[DEBUG] create_word_document called for: {unite.get('titre_unite')}")
    
    # Charger le template (chemin relatif pour Vercel)
    template_path = os.path.join(os.path.dirname(__file__), 'public', 'Unité PEI.docx')
    print(f"[DEBUG] Template path: {template_path}")
    print(f"[DEBUG] Template exists: {os.path.exists(template_path)}")
    
    doc = Document(template_path)
    print(f"[DEBUG] Template loaded, {len(doc.tables)} tables found")
    
    # Remplacer les placeholders dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Remplacer les placeholders
                text = cell.text
                text = text.replace('{enseignant}', enseignant)
                text = text.replace('{groupe_matiere}', matiere_data['nom'])
                text = text.replace('{titre_unite}', unite.get('titre_unite', ''))
                text = text.replace('{annee_pei}', str(annee_pei))
                text = text.replace('{duree}', str(unite.get('duree', '')))
                text = text.replace('{concept_cle}', unite.get('concept_cle', ''))
                text = text.replace('{concepts_connexes}', ', '.join(unite.get('concepts_connexes', [])))
                text = text.replace('{contexte_mondial}', unite.get('contexte_mondial', ''))
                text = text.replace('{enonce_de_recherche}', unite.get('enonce_recherche', ''))
                text = text.replace('{questions_factuelles}', '\n'.join(unite.get('questions_factuelles', [])))
                text = text.replace('{questions_conceptuelles}', '\n'.join(unite.get('questions_conceptuelles', [])))
                text = text.replace('{questions_debat}', '\n'.join(unite.get('questions_debat', [])))
                
                # Formater les objectifs spécifiques
                objectifs_text = format_objectifs_specifiques(
                    unite.get('objectifs_specifiques', []),
                    matiere_data,
                    annee_pei
                )
                text = text.replace('{objectifs_specifiques}', objectifs_text)
                
                # Clear and rewrite cell content
                cell.text = text
    
    # Sauvegarder le document
    # Sur Vercel, utiliser /tmp pour l'écriture (système de fichiers en lecture seule ailleurs)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Unite_PEI_{matiere_data['nom']}_{timestamp}.docx"
    
    # Utiliser /tmp sur Vercel, generated_units/ en local
    if os.path.exists('/tmp'):
        # Environnement Vercel
        filepath = os.path.join('/tmp', filename)
        print(f"[DEBUG] Using Vercel /tmp directory: {filepath}")
    else:
        # Environnement local
        filepath = os.path.join(app.config['GENERATED_UNITS_FOLDER'], filename)
        os.makedirs(app.config['GENERATED_UNITS_FOLDER'], exist_ok=True)
        print(f"[DEBUG] Using local directory: {filepath}")
    
    doc.save(filepath)
    print(f"[DEBUG] Document saved to: {filepath}")
    
    return filename

def format_objectifs_specifiques(objectifs_ids, matiere_data, annee_pei):
    """Formate les objectifs spécifiques pour le document"""
    objectifs_dict = matiere_data.get('objectifs', {}).get(annee_pei, {})
    
    formatted = []
    current_criterion = None
    
    for obj_id in objectifs_ids:
        # Parse "A.i" -> criterion="A", number="i"
        parts = obj_id.split('.')
        if len(parts) == 2:
            criterion, number = parts
            
            if criterion != current_criterion:
                criterion_names = {
                    'A': 'Recherche et analyse' if 'design' in str(matiere_data) else 'Analyse',
                    'B': 'Développement des idées' if 'design' in str(matiere_data) else 'Organisation',
                    'C': 'Création de la solution' if 'design' in str(matiere_data) else 'Production de texte',
                    'D': 'Évaluation' if 'design' in str(matiere_data) else 'Utilisation de la langue'
                }
                formatted.append(f"\n{criterion}: {criterion_names.get(criterion, criterion)}")
                current_criterion = criterion
            
            # Trouver le texte de l'objectif
            if criterion in objectifs_dict:
                for obj in objectifs_dict[criterion]:
                    if obj.startswith(f"{number}."):
                        formatted.append(obj)
                        break
    
    return '\n'.join(formatted)

@app.route('/download/<filename>')
def download_file(filename):
    """Télécharge un fichier généré"""
    try:
        # Chercher d'abord dans /tmp (Vercel), puis dans generated_units/ (local)
        tmp_path = os.path.join('/tmp', filename)
        local_path = os.path.join(app.config['GENERATED_UNITS_FOLDER'], filename)
        
        if os.path.exists(tmp_path):
            filepath = tmp_path
            print(f"[DEBUG] Serving file from /tmp: {filename}")
        elif os.path.exists(local_path):
            filepath = local_path
            print(f"[DEBUG] Serving file from local directory: {filename}")
        else:
            print(f"[ERROR] File not found: {filename}")
            print(f"[ERROR] Checked paths: {tmp_path}, {local_path}")
            return jsonify({"error": f"File not found: {filename}"}), 404
        
        return send_file(filepath, as_attachment=True)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[ERROR] Exception in download_file: {str(e)}")
        print(f"[ERROR] Traceback: {error_details}")
        return jsonify({"error": str(e), "details": error_details}), 404

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
