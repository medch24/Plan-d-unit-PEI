#!/usr/bin/env python3
"""
Create clean Word templates with proper placeholders that won't be fragmented
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_placeholder(paragraph, text, bold=False, size=None):
    """Add placeholder text to paragraph"""
    run = paragraph.add_run(text)
    if bold:
        run.font.bold = True
    if size:
        run.font.size = Pt(size)
    return run

def create_plan_template():
    """Create Plan d'Unité template with proper placeholders"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder(title, 'Plan de travail de l\'unité PEI', bold=True, size=16)
    
    # Info table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Row 1: Enseignant
    table.rows[0].cells[0].text = 'Enseignant(s)'
    add_placeholder(table.rows[0].cells[1].paragraphs[0], '{enseignant}')
    
    # Row 2: Titre
    table.rows[1].cells[0].text = 'Titre de l\'unité'
    add_placeholder(table.rows[1].cells[1].paragraphs[0], '{titre_unite}')
    
    # Row 3: Matière
    table.rows[2].cells[0].text = 'Groupe de matières et discipline'
    add_placeholder(table.rows[2].cells[1].paragraphs[0], '{groupe_matiere}')
    
    # Row 4: Année
    table.rows[3].cells[0].text = 'Année du PEI'
    add_placeholder(table.rows[3].cells[1].paragraphs[0], '{annee_pei}')
    
    # Row 5: Durée
    table.rows[4].cells[0].text = 'Durée de l\'unité (heures)'
    add_placeholder(table.rows[4].cells[1].paragraphs[0], '{duree}')
    
    # Row 6: Merged header for research section
    cells = table.rows[5].cells
    cells[0].merge(cells[1])
    add_placeholder(cells[0].paragraphs[0], 'Recherche : définition de l\'objectif de l\'unité', bold=True, size=12)
    
    doc.add_paragraph()
    
    # Research table
    research_table = doc.add_table(rows=5, cols=2)
    research_table.style = 'Table Grid'
    
    research_table.rows[0].cells[0].text = 'Concept clé'
    add_placeholder(research_table.rows[0].cells[1].paragraphs[0], '{concept_cle}')
    
    research_table.rows[1].cells[0].text = 'Concept(s) connexe(s)'
    add_placeholder(research_table.rows[1].cells[1].paragraphs[0], '{concepts_connexes}')
    
    research_table.rows[2].cells[0].text = 'Contexte mondial'
    add_placeholder(research_table.rows[2].cells[1].paragraphs[0], '{contexte_mondial}')
    
    research_table.rows[3].cells[0].text = 'Énoncé de recherche'
    add_placeholder(research_table.rows[3].cells[1].paragraphs[0], '{enonce_de_recherche}')
    
    research_table.rows[4].cells[0].text = 'Questions de recherche'
    questions_cell = research_table.rows[4].cells[1]
    questions_cell.paragraphs[0].add_run('Factuelle(s): ').bold = True
    add_placeholder(questions_cell.paragraphs[0], '{questions_factuelles}')
    questions_cell.add_paragraph()
    questions_cell.paragraphs[1].add_run('Conceptuelle(s): ').bold = True
    add_placeholder(questions_cell.paragraphs[1], '{questions_conceptuelles}')
    questions_cell.add_paragraph()
    questions_cell.paragraphs[2].add_run('Invitant au débat: ').bold = True
    add_placeholder(questions_cell.paragraphs[2], '{questions_debat}')
    
    doc.add_paragraph()
    
    # Objectives section
    objectives_table = doc.add_table(rows=1, cols=1)
    objectives_table.style = 'Table Grid'
    objectives_table.rows[0].cells[0].text = 'Objectifs spécifiques'
    doc.add_paragraph()
    add_placeholder(doc.add_paragraph(), '{objectifs_specifiques}')
    
    doc.add_paragraph()
    
    # Evaluation sommative
    doc.add_heading('Évaluation sommative', level=2)
    add_placeholder(doc.add_paragraph(), '{evaluation_sommative}')
    
    doc.add_paragraph()
    
    # Approches de l'apprentissage
    doc.add_heading('Approches de l\'apprentissage', level=2)
    add_placeholder(doc.add_paragraph(), '{approches_apprentissage}')
    
    doc.add_paragraph()
    
    # Contenu
    doc.add_heading('Contenu et processus d\'apprentissage', level=2)
    add_placeholder(doc.add_paragraph(), '{contenu}')
    
    doc.add_paragraph()
    
    # Ressources
    doc.add_heading('Ressources', level=2)
    add_placeholder(doc.add_paragraph(), '{ressources}')
    
    doc.add_paragraph()
    
    # Différenciation
    doc.add_heading('Différenciation', level=2)
    add_placeholder(doc.add_paragraph(), '{differenciation}')
    
    doc.add_paragraph()
    
    # Évaluation formative
    doc.add_heading('Évaluation formative', level=2)
    add_placeholder(doc.add_paragraph(), '{evaluation_formative}')
    
    doc.add_paragraph()
    
    # Réflexions
    doc.add_heading('Réflexion', level=2)
    doc.add_heading('Avant l\'enseignement', level=3)
    add_placeholder(doc.add_paragraph(), '{reflexion_avant}')
    doc.add_paragraph()
    doc.add_heading('Pendant l\'enseignement', level=3)
    add_placeholder(doc.add_paragraph(), '{reflexion_pendant}')
    doc.add_paragraph()
    doc.add_heading('Après l\'enseignement', level=3)
    add_placeholder(doc.add_paragraph(), '{reflexion_apres}')
    
    # Save
    doc.save('/home/user/webapp/templates/Plan_CLEAN_TEMPLATE.docx')
    print("✅ Plan template created: Plan_CLEAN_TEMPLATE.docx")
    print("   Placeholders: 23 total")

def create_eval_template():
    """Create Evaluation template with proper placeholders"""
    doc = Document()
    
    # Title
    title = doc.add_paragraph('Évaluation critériée – ', style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder(title, '{groupe_matiere}', bold=True)
    
    # Unit info
    p = doc.add_paragraph('Unité: ')
    add_placeholder(p, '{titre_unite}')
    
    p = doc.add_paragraph('Énoncé de recherche: ')
    add_placeholder(p, '{enonce_de_recherche}')
    
    p = doc.add_paragraph('Année du PEI: ')
    add_placeholder(p, '{annee_pei}')
    
    doc.add_paragraph()
    
    # Criterion header
    criterion_header = doc.add_heading(level=2)
    add_placeholder(criterion_header, 'Critère {lettre_critere} : {nom_objectif_specifique}', bold=True)
    
    doc.add_paragraph()
    
    # Objectifs spécifiques
    doc.add_heading('Objectifs spécifiques évalués', level=3)
    add_placeholder(doc.add_paragraph(), '{objectifs_specifiques}')
    
    doc.add_paragraph()
    
    # Exercices section
    doc.add_heading('Exercices d\'évaluation', level=2)
    add_placeholder(doc.add_paragraph(), '{exercices}')
    
    doc.add_paragraph()
    
    # Grille d'évaluation
    doc.add_heading('Grille d\'évaluation', level=2)
    
    # Descriptors table
    desc_table = doc.add_table(rows=5, cols=2)
    desc_table.style = 'Table Grid'
    
    desc_table.rows[0].cells[0].text = 'Niveaux'
    desc_table.rows[0].cells[1].text = 'Descripteurs'
    
    desc_table.rows[1].cells[0].text = '1-2'
    add_placeholder(desc_table.rows[1].cells[1].paragraphs[0], '{descripteur_1_2}')
    
    desc_table.rows[2].cells[0].text = '3-4'
    add_placeholder(desc_table.rows[2].cells[1].paragraphs[0], '{descripteur_3_4}')
    
    desc_table.rows[3].cells[0].text = '5-6'
    add_placeholder(desc_table.rows[3].cells[1].paragraphs[0], '{descripteur_5_6}')
    
    desc_table.rows[4].cells[0].text = '7-8'
    add_placeholder(desc_table.rows[4].cells[1].paragraphs[0], '{descripteur_7_8}')
    
    doc.add_paragraph()
    
    # Student workspace
    doc.add_heading('Espace de travail pour l\'élève', level=3)
    doc.add_paragraph('1. .............................................................')
    doc.add_paragraph('2. .............................................................')
    doc.add_paragraph('3. .............................................................')
    doc.add_paragraph('[Espace pour insérer une image/ressource]')
    
    # Save
    doc.save('/home/user/webapp/templates/Eval_CLEAN_TEMPLATE.docx')
    print("✅ Eval template created: Eval_CLEAN_TEMPLATE.docx")
    print("   Placeholders: 12 total")

if __name__ == '__main__':
    print("🔧 Creating clean Word templates...")
    print("="*60)
    create_plan_template()
    create_eval_template()
    print("="*60)
    print("✅ Templates created successfully!")
    print("\nPlaceholders summary:")
    print("\nPlan Template:")
    print("  - enseignant, titre_unite, groupe_matiere, annee_pei, duree")
    print("  - concept_cle, concepts_connexes, contexte_mondial")
    print("  - enonce_de_recherche")
    print("  - questions_factuelles, questions_conceptuelles, questions_debat")
    print("  - objectifs_specifiques")
    print("  - evaluation_sommative, approches_apprentissage")
    print("  - contenu, ressources, differenciation")
    print("  - evaluation_formative")
    print("  - reflexion_avant, reflexion_pendant, reflexion_apres")
    print("\nEval Template:")
    print("  - groupe_matiere, titre_unite, enonce_de_recherche, annee_pei")
    print("  - lettre_critere, nom_objectif_specifique")
    print("  - objectifs_specifiques, exercices")
    print("  - descripteur_1_2, descripteur_3_4, descripteur_5_6, descripteur_7_8")
